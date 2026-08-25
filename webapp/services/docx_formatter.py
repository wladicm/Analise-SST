"""Formatação de documentos .docx: quebras de página, alinhamento, tabelas
e, opcionalmente, padrão ABNT NBR 14724 (margens, fonte, espaçamento,
recuo de primeira linha e numeração de página)."""
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def formatar(caminho_origem: str, caminho_destino: str, opcoes: dict):
    doc = Document(caminho_origem)

    quebra_pagina = opcoes.get("quebra_pagina_titulos", True)
    justificar = opcoes.get("justificar_corpo", True)
    centralizar_titulos = opcoes.get("centralizar_titulos", True)
    padronizar_tabelas = opcoes.get("padronizar_tabelas", True)
    aplicar_abnt = opcoes.get("aplicar_abnt", False)

    primeiro_titulo_processado = False

    for paragrafo in doc.paragraphs:
        estilo = paragrafo.style.name if paragrafo.style else ""
        eh_titulo = "Heading 1" in estilo or "Título 1" in estilo or "Heading1" in estilo

        if eh_titulo:
            if quebra_pagina and primeiro_titulo_processado:
                paragrafo.paragraph_format.page_break_before = True
            if centralizar_titulos:
                paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            primeiro_titulo_processado = True
        elif justificar and paragrafo.text.strip():
            paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if aplicar_abnt:
            _aplicar_estilo_abnt(paragrafo, eh_titulo)

    if padronizar_tabelas:
        for tabela in doc.tables:
            _padronizar_tabela(tabela)

    if aplicar_abnt:
        _aplicar_margens_abnt(doc)
        _adicionar_numeracao_pagina(doc)

    doc.save(caminho_destino)


def _aplicar_estilo_abnt(paragrafo, eh_titulo: bool):
    fmt = paragrafo.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    if not eh_titulo:
        fmt.first_line_indent = Cm(1.25)

    for run in paragrafo.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        # Garante que a fonte também se aplique a textos do Extremo Oriente,
        # necessário em alguns arquivos gerados por outros editores.
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def _aplicar_margens_abnt(doc):
    for secao in doc.sections:
        secao.top_margin = Cm(3)
        secao.left_margin = Cm(3)
        secao.bottom_margin = Cm(2)
        secao.right_margin = Cm(2)


def _adicionar_numeracao_pagina(doc):
    """Insere campo PAGE no cabeçalho, alinhado à direita (padrão ABNT:
    número no canto superior direito, a 2cm da borda superior)."""
    for secao in doc.sections:
        header = secao.header
        header.is_linked_to_previous = False
        paragrafo = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragrafo.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in list(paragrafo.runs):
            run.clear()

        run = paragrafo.add_run()
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)

        fld_begin = OxmlElement('w:fldChar')
        fld_begin.set(qn('w:fldCharType'), 'begin')

        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = ' PAGE '

        fld_separate = OxmlElement('w:fldChar')
        fld_separate.set(qn('w:fldCharType'), 'separate')

        fld_end = OxmlElement('w:fldChar')
        fld_end.set(qn('w:fldCharType'), 'end')

        run._r.append(fld_begin)
        run._r.append(instr_text)
        run._r.append(fld_separate)
        run._r.append(fld_end)


def _padronizar_tabela(tabela):
    tabela.autofit = False
    num_colunas = len(tabela.columns)
    if num_colunas == 0:
        return

    largura_coluna = Cm(16) / num_colunas  # ~16cm de largura útil (A4 com margens ABNT)

    for i, linha in enumerate(tabela.rows):
        for celula in linha.cells:
            celula.width = largura_coluna
            for paragrafo in celula.paragraphs:
                paragrafo.alignment = (
                    WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                )
            _definir_bordas_celula(celula)


def _definir_bordas_celula(celula):
    tc_pr = celula._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for lado in ('top', 'left', 'bottom', 'right'):
        elemento = OxmlElement(f'w:{lado}')
        elemento.set(qn('w:val'), 'single')
        elemento.set(qn('w:sz'), '4')
        elemento.set(qn('w:color'), '000000')
        tc_borders.append(elemento)
    tc_pr.append(tc_borders)
